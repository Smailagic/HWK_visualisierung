from docx import Document
from docx.shared import RGBColor

def extrahiere_rote_woerter(input_datei, output_datei):
    # Word-Dokument laden
    doc = Document(input_datei)
    
    # Liste zum Speichern der roten Wörter
    rote_woerter = []

    # Durch alle Absätze und Runs im Dokument gehen
    for para in doc.paragraphs:
        for run in para.runs:
            # Prüfen, ob der Text rot ist (RGB-Wert: 255 für rot)
            if run.font.color and run.font.color.rgb == RGBColor(255, 0, 0):  # Rot: RGB(255, 0, 0)
                # Das Wort aus dem Run extrahieren
                rote_woerter.append(run.text.strip())
    
    # Neues Dokument erstellen und die gefundenen roten Wörter hinzufügen
    doc_rot = Document()
    doc_rot.add_heading('Rote Wörter', 0)
    for wort in rote_woerter:
        doc_rot.add_paragraph(wort)

    # Speichern der neuen Datei
    doc_rot.save(output_datei)

# Verwendung des Skripts
input_datei = 'Präpositionen.docx'  # Der Name des Eingabedokuments
output_datei = 'ROT.docx'  # Der Name der Ausgabedatei

extrahiere_rote_woerter(input_datei, output_datei)

